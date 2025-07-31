import docx
import json
from typing import List, Dict

class DocumentProcessor:
    def extract_text(self, filepath: str) -> str:
        """Extract text from a document file."""
        if filepath.endswith('.docx'):
            return self._extract_from_docx(filepath)
        elif filepath.endswith('.txt'):
            return self._extract_from_txt(filepath)
        else:
            raise ValueError(f"Unsupported file type: {filepath}")
    
    def _extract_from_docx(self, filepath: str) -> str:
        """Extract text from a DOCX file."""
        doc = docx.Document(filepath)
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(' | '.join(row_text))
        
        return '\n'.join(full_text)
    
    def _extract_from_txt(self, filepath: str) -> str:
        """Extract text from a TXT file."""
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    
    def generate_presentation_plan(self, text_content: str, llm) -> List[Dict]:
        """Generate a presentation plan from document text using LLM."""
        prompt = f"""You are a presentation structuring expert. Analyze the following text from a document and propose a slide-by-slide plan in JSON format. 

IMPORTANT INSTRUCTIONS:
1. Create logical groupings of content into slides
2. Each slide should have a clear title
3. Identify if content should be bullet points or a table
4. For tables, structure the data properly
5. Keep bullet points concise and clear
6. Aim for 5-15 slides depending on content volume

DOCUMENT TEXT:
{text_content}

OUTPUT FORMAT:
Return ONLY a JSON array with this structure:
[
  {{
    "slide": 1,
    "title": "Clear slide title",
    "content": ["Bullet point 1", "Bullet point 2", "..."]
  }},
  {{
    "slide": 2,
    "title": "Another slide title",
    "content_type": "table",
    "data": [["Header 1", "Header 2"], ["Row 1 Col 1", "Row 1 Col 2"], ...]
  }}
]

Ensure the JSON is valid and properly formatted.
"""
        
        response = llm.generate_response(prompt)
        
        # Extract JSON from response
        try:
            # Clean up the response - remove markdown code blocks if present
            cleaned_response = response.strip()
            
            # Remove markdown code block markers
            if cleaned_response.startswith('```json'):
                cleaned_response = cleaned_response[7:]  # Remove ```json
            elif cleaned_response.startswith('```'):
                cleaned_response = cleaned_response[3:]  # Remove ```
                
            if cleaned_response.endswith('```'):
                cleaned_response = cleaned_response[:-3]  # Remove ending ```
            
            # Try to find JSON in the response
            json_start = cleaned_response.find('[')
            json_end = cleaned_response.rfind(']') + 1
            
            if json_start != -1 and json_end > json_start:
                json_str = cleaned_response[json_start:json_end]
                plan = json.loads(json_str)
                return plan
            else:
                # Try parsing the whole cleaned response
                plan = json.loads(cleaned_response.strip())
                return plan
                
        except json.JSONDecodeError as e:
            # Log the error for debugging
            print(f"JSON parsing error: {e}")
            print(f"Response was: {response[:200]}...")
            # Fallback to a simple plan if JSON parsing fails
            pass
        
        # Fallback plan
        return [{
            "slide": 1,
            "title": "Presentation",
            "content": ["Content could not be automatically structured", "Please edit manually"]
        }]